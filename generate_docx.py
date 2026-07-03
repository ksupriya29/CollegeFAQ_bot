"""
Script to generate the college_info.docx knowledge base document.
"""

from docx import Document


def generate_college_info():
    doc = Document()
    
    # Title
    doc.add_heading("BVRIT Hyderabad College of Engineering for Women", 0)
    doc.add_paragraph("Official Information Document")
    doc.add_paragraph("")
    
    # ===== Section 1: About BVRIT =====
    doc.add_heading("1. About BVRIT", level=1)
    doc.add_heading("1.1 History", level=2)
    doc.add_paragraph(
        "BVRIT Hyderabad College of Engineering for Women was established in 1997 "
        "by the Sri Vishnu Educational Society. The college is located in Bachupally, "
        "Hyderabad, Telangana. It is a premier institution dedicated to women's education "
        "in engineering and technology, with over 3,000 students and 200+ faculty members."
    )
    
    doc.add_heading("1.2 Vision", level=2)
    doc.add_paragraph(
        "To emerge as a premier institution for women's education in engineering and "
        "technology, fostering academic excellence, research, and holistic development."
    )
    
    doc.add_heading("1.3 Mission", level=2)
    doc.add_paragraph(
        "To provide quality education through innovative teaching methods, state-of-the-art "
        "infrastructure, and industry collaboration. To empower women engineers with "
        "technical competence, leadership skills, and ethical values."
    )
    
    doc.add_heading("1.4 Accreditations", level=2)
    doc.add_paragraph(
        "BVRIT Hyderabad is NAAC Accredited with 'A' Grade. The B.Tech programs in CSE, "
        "ECE, EEE, and ME are NBA Accredited. The college is an Autonomous Institution "
        "affiliated to JNTU Hyderabad and approved by AICTE, New Delhi. It is also "
        "ISO 9001:2015 Certified."
    )
    
    doc.add_heading("1.5 Rankings", level=2)
    doc.add_paragraph(
        "BVRIT is ranked among the top engineering colleges in Telangana. It is recognized "
        "for academic excellence, placement record, and campus facilities."
    )
    
    # ===== Section 2: Departments =====
    doc.add_heading("2. Departments", level=1)
    doc.add_paragraph(
        "BVRIT offers the following B.Tech programs (4 years duration):\n\n"
        "1. Computer Science & Engineering (CSE) - 120 seats\n"
        "2. Electronics & Communication Engineering (ECE) - 120 seats\n"
        "3. Electrical & Electronics Engineering (EEE) - 60 seats\n"
        "4. Mechanical Engineering (ME) - 60 seats\n"
        "5. Artificial Intelligence & Machine Learning (AI & ML) - 60 seats\n"
        "6. Computer Science & Engineering (Data Science) - 60 seats\n"
        "7. Information Technology (IT) - 60 seats\n\n"
        "Total seats across all B.Tech programs: Approximately 540+ seats."
    )
    
    doc.add_heading("2.1 M.Tech Programs", level=2)
    doc.add_paragraph(
        "M.Tech in Computer Science & Engineering\n"
        "M.Tech in VLSI System Design\n"
        "M.Tech in Power Electronics\n\n"
        "All M.Tech programs are 2 years duration."
    )
    
    doc.add_heading("2.2 Faculty", level=2)
    doc.add_paragraph(
        "The college has 200+ highly qualified faculty members, many with PhDs. "
        "Faculty are involved in research, publishing papers in international journals, "
        "and guiding student projects."
    )
    
    # ===== Section 3: Admissions =====
    doc.add_heading("3. Admissions", level=1)
    doc.add_heading("3.1 Eligibility for B.Tech", level=2)
    doc.add_paragraph(
        "Minimum 45% marks in 10+2 (or equivalent) with Physics, Chemistry, and Mathematics. "
        "Valid TS EAMCET or JEE Main score is required. For Management quota, direct admission "
        "is available based on 10+2 marks."
    )
    
    doc.add_heading("3.2 Eligibility for M.Tech", level=2)
    doc.add_paragraph(
        "B.Tech or BE in relevant discipline with minimum 55% marks. Valid GATE or TS PGECET "
        "score is required."
    )
    
    doc.add_heading("3.3 Application Process", level=2)
    doc.add_paragraph(
        "1. Visit the official website bvrithyderabad.edu.in\n"
        "2. Fill the online application form\n"
        "3. Upload required documents (marksheets, EAMCET rank card, ID proof, photos)\n"
        "4. Pay application fee of Rs. 1,000\n"
        "5. Attend counseling (EAMCET) or direct interview (Management quota)"
    )
    
    doc.add_heading("3.4 Required Documents", level=2)
    doc.add_paragraph(
        "10th & 12th Marksheets, TS EAMCET Rank Card or JEE Scorecard, Transfer Certificate (TC), "
        "Migration Certificate (if applicable), Caste Certificate (if applicable), "
        "Income Certificate (for fee reimbursement), Passport size photographs, Aadhar Card."
    )
    
    doc.add_heading("3.5 TS EAMCET", level=2)
    doc.add_paragraph(
        "BVRIT College Code: BVRT. Previous year cutoff ranks (approximate): CSE - Rank under 10,000, "
        "ECE - Rank under 15,000, EEE - Rank under 25,000, ME - Rank under 35,000, "
        "AI & ML - Rank under 12,000, Data Science - Rank under 14,000. "
        "Admissions open from May to July each year."
    )
    
    # ===== Section 4: Fee Structure =====
    doc.add_heading("4. Fee Structure (2024-25)", level=1)
    doc.add_heading("4.1 B.Tech Tuition Fees (Per Year)", level=2)
    doc.add_paragraph(
        "CSE, AI & ML, Data Science, IT: Rs. 1,10,000\n"
        "ECE, EEE: Rs. 90,000\n"
        "Mechanical Engineering: Rs. 80,000"
    )
    
    doc.add_heading("4.2 M.Tech Tuition Fees (Per Year)", level=2)
    doc.add_paragraph("All Specializations: Rs. 70,000 per year.")
    
    doc.add_heading("4.3 Additional Fees", level=2)
    doc.add_paragraph(
        "Application Fee: Rs. 1,000 (one-time)\n"
        "Library Fee: Rs. 3,000 per year\n"
        "Laboratory Fee: Rs. 5,000 per year\n"
        "Sports & Cultural Fee: Rs. 2,000 per year\n"
        "Hostel Fee: Rs. 50,000 to Rs. 70,000 per year (including mess)\n"
        "Transport Fee: Rs. 12,000 to Rs. 18,000 per year"
    )
    
    doc.add_heading("4.4 Hostel Fee Breakup (Per Year)", level=2)
    doc.add_paragraph(
        "Single Room: Rs. 70,000\n"
        "Double Sharing: Rs. 55,000\n"
        "Triple Sharing: Rs. 45,000\n"
        "Mess Fee (additional): Rs. 30,000 per year"
    )
    
    doc.add_heading("4.5 Payment Options", level=2)
    doc.add_paragraph(
        "Online (Net Banking, Card, UPI), Demand Draft, or Bank Transfer."
    )
    
    # ===== Section 5: Scholarships =====
    doc.add_heading("5. Scholarships", level=1)
    doc.add_heading("5.1 TS EAMCET Rank Holders", level=2)
    doc.add_paragraph(
        "Rank 1-1000: 100% tuition fee waiver\n"
        "Rank 1001-5000: 50% tuition fee waiver\n"
        "Rank 5001-10000: 25% tuition fee waiver"
    )
    
    doc.add_heading("5.2 Merit-Based Scholarships", level=2)
    doc.add_paragraph(
        "Semester toppers: Rs. 10,000 cash prize plus certificate\n"
        "Overall CGPA above 9.0: 25% fee concession for next semester\n"
        "100% attendance in a semester: Rs. 5,000 incentive"
    )
    
    doc.add_heading("5.3 Government Schemes", level=2)
    doc.add_paragraph(
        "TS E-Pass (Fee Reimbursement for SC/ST/BC/EBC students), "
        "National Scholarship Portal (NSP) schemes, Minority Community Scholarships, "
        "Single Girl Child Scholarship."
    )
    
    doc.add_heading("5.4 Sports Quota", level=2)
    doc.add_paragraph("National/State level players: Up to 50% fee waiver.")
    
    # ===== Section 6: Placements =====
    doc.add_heading("6. Placements", level=1)
    doc.add_heading("6.1 Placement Statistics (2023-24)", level=2)
    doc.add_paragraph(
        "Overall Placement Rate: 90%+\n"
        "Highest Package: Rs. 24 LPA (CSE - Amazon)\n"
        "Average Package: Rs. 6.5 LPA\n"
        "Number of Recruiters: 100+\n"
        "Multiple offers: 30% of students got 2+ offers"
    )
    
    doc.add_heading("6.2 Top Recruiting Companies", level=2)
    doc.add_paragraph(
        "Tech Giants: Google, Microsoft, Amazon, Meta\n"
        "IT Services: Infosys, TCS, Wipro, HCL, Tech Mahindra, Accenture\n"
        "Consulting: Deloitte, KPMG, EY, PwC\n"
        "Finance: Goldman Sachs, JP Morgan, Barclays\n"
        "Product: Flipkart, Uber, Ola, Swiggy, Zomato\n"
        "Core: L&T, Siemens, Bosch, Tata Motors, Mahindra"
    )
    
    doc.add_heading("6.3 Department-wise Average Package", level=2)
    doc.add_paragraph(
        "CSE / AI / Data Science: Rs. 8-12 LPA\n"
        "ECE: Rs. 5-8 LPA\n"
        "EEE: Rs. 4-7 LPA\n"
        "ME: Rs. 4-6 LPA"
    )
    
    doc.add_heading("6.4 Placement Training", level=2)
    doc.add_paragraph(
        "Resume building and LinkedIn profile optimization. "
        "Mock interviews and aptitude tests. "
        "Soft skills and communication training. "
        "Industry guest lectures and workshops. "
        "Internship opportunities with partner companies. "
        "Placement season runs from August to March each year."
    )
    
    # ===== Section 7: Campus & Facilities =====
    doc.add_heading("7. Campus & Facilities", level=1)
    doc.add_heading("7.1 Academic Facilities", level=2)
    doc.add_paragraph(
        "Smart classrooms with projectors and audio systems. "
        "Well-equipped laboratories for each department. "
        "Computer labs with latest hardware and software. "
        "Central library with digital resources. "
        "Seminar halls with 200+ seating capacity. "
        "Open-air auditorium for events."
    )
    
    doc.add_heading("7.2 Library", level=2)
    doc.add_paragraph(
        "40,000+ books across all disciplines. "
        "150+ national and international journals. "
        "8,000+ e-books and e-journals. "
        "Access to IEEE, Springer, Elsevier digital libraries. "
        "Previous year question papers and project reports. "
        "Library timings: Monday-Friday 8 AM to 7 PM, Saturday 9 AM to 4 PM. "
        "During exams: Extended hours till 9 PM."
    )
    
    doc.add_heading("7.3 Hostel", level=2)
    doc.add_paragraph(
        "Separate hostel accommodation for girls within the campus. "
        "Well-furnished rooms available in Single, Double, and Triple sharing. "
        "24/7 Wi-Fi connectivity. Round-the-clock security with CCTV surveillance. "
        "Purified drinking water (RO). Solar-powered hot water. "
        "Common room with TV and indoor games. "
        "Study rooms with adequate lighting. Laundry facilities. "
        "Power backup (Generator). "
        "Vegetarian mess with nutritious meals. "
        "Special menu on weekends and festivals. "
        "Canteen for additional snacks."
    )
    
    doc.add_heading("7.4 Sports & Fitness", level=2)
    doc.add_paragraph(
        "Outdoor: Basketball court, Volleyball court, Badminton court, "
        "400m athletic track, Multi-purpose sports field. "
        "Indoor: Table tennis room, Chess and carrom room, Indoor games area. "
        "Fitness: Gym with modern equipment, Yoga and meditation room, "
        "Zumba and aerobics sessions. "
        "Annual Sports Meet in February."
    )
    
    doc.add_heading("7.5 Support Facilities", level=2)
    doc.add_paragraph(
        "24/7 Wi-Fi across campus. Medical clinic with resident doctor. "
        "Cafeteria and food court. Bank and ATM on campus. "
        "Student parking. Stationery store. Girls common room. "
        "Wheelchair ramps and elevators for accessibility. "
        "Green campus with solar panels, rainwater harvesting, organic garden."
    )
    
    doc.add_heading("7.6 Transport", level=2)
    doc.add_paragraph(
        "30+ buses covering major routes in Hyderabad including Kukatpally, KPHB, "
        "Miyapur, Chandanagar, Madhapur, Gachibowli, Hitech City, Ameerpet, "
        "Secunderabad, Dilsukhnagar, LB Nagar, Uppal. "
        "Morning pickup: 6:30 AM to 8:30 AM. Evening drop: 4:00 PM to 6:00 PM. "
        "Transport fee: Rs. 12,000 to Rs. 18,000 per year depending on route."
    )
    
    # ===== Section 8: Faculty =====
    doc.add_heading("8. Faculty", level=1)
    doc.add_paragraph(
        "BVRIT has 200+ highly qualified faculty members. Many faculty members hold "
        "PhD degrees from prestigious universities. Faculty are actively involved in "
        "research with 300+ research papers published in international journals in the "
        "last 3 years. 25+ patents have been filed with 10 granted. "
        "20+ funded research projects are ongoing. Faculty research centers include: "
        "Center for AI and Machine Learning Research, VLSI Design and Embedded Systems Lab, "
        "Renewable Energy Research Lab, and Data Analytics Center."
    )
    
    # ===== Section 9: Student Life =====
    doc.add_heading("9. Student Life", level=1)
    doc.add_heading("9.1 Technical Clubs", level=2)
    doc.add_paragraph(
        "CodeCraft (Coding and Competitive Programming), Robotics Club, "
        "AI/ML Enthusiasts Group, Cybersecurity Club, Web Development Society."
    )
    
    doc.add_heading("9.2 Cultural Clubs", level=2)
    doc.add_paragraph(
        "Dance Club (Classical, Western), Music Club (Vocal, Instrumental), "
        "Drama and Theatre Society, Fine Arts and Photography Club, "
        "Literary and Debate Club."
    )
    
    doc.add_heading("9.3 Special Interest Groups", level=2)
    doc.add_paragraph(
        "Women Empowerment Cell, Entrepreneurship Cell (E-Cell), "
        "NSS (National Service Scheme), Environment Club, Yoga and Meditation Club."
    )
    
    doc.add_heading("9.4 Annual Events", level=2)
    doc.add_paragraph(
        "Technovanza (Technical Fest - March), Cultural Fest (October), "
        "Sports Meet (February), Women's Day Celebrations (March 8)."
    )
    
    # ===== Section 10: Examination =====
    doc.add_heading("10. Examination & Grading", level=1)
    doc.add_heading("10.1 Examination Pattern", level=2)
    doc.add_paragraph(
        "Two semesters per year (Odd: July-December, Even: January-June). "
        "Internal Assessment: 40% (assignments, quizzes, mid-term exams). "
        "End Semester Exam: 60% (conducted by JNTU Hyderabad). "
        "Practical exams for lab-based subjects. Project work with viva voce."
    )
    
    doc.add_heading("10.2 Grading System", level=2)
    doc.add_paragraph(
        "O Grade: 90%+ (10 points)\n"
        "A+ Grade: 80-89% (9 points)\n"
        "A Grade: 70-79% (8 points)\n"
        "B+ Grade: 60-69% (7 points)\n"
        "B Grade: 55-59% (6 points)\n"
        "C Grade: 50-54% (5 points) - Pass\n"
        "F Grade: Below 50% - Fail"
    )
    
    doc.add_heading("10.3 Attendance Policy", level=2)
    doc.add_paragraph(
        "Minimum 75% attendance in each subject to appear for exams. "
        "Attendance below 65%: Not eligible for semester exams. "
        "Attendance 65-74%: Can apply for condonation with medical proof."
    )
    
    # ===== Section 11: Contact =====
    doc.add_heading("11. Contact Information", level=1)
    doc.add_paragraph(
        "Address:\n"
        "BVRIT Hyderabad College of Engineering for Women\n"
        "Bachupally, Nizampet Road\n"
        "Hyderabad - 500090, Telangana, India\n\n"
        "Phone:\n"
        "Admissions: +91-40-2304-2777\n"
        "General Enquiry: +91-40-2304-2778\n"
        "Helpline: 1800-123-4567 (Toll Free)\n\n"
        "Email:\n"
        "Admissions: admissions@bvrithyderabad.edu.in\n"
        "General: info@bvrithyderabad.edu.in\n"
        "Placements: placements@bvrithyderabad.edu.in\n\n"
        "Website: bvrithyderabad.edu.in\n\n"
        "Office Hours:\n"
        "Monday - Friday: 9:00 AM - 5:00 PM\n"
        "Saturday: 9:00 AM - 1:00 PM\n"
        "Sunday: Closed"
    )
    
    doc.add_heading("11.1 Anti-Ragging & Grievance", level=2)
    doc.add_paragraph(
        "Ragging is strictly prohibited as per UGC guidelines. "
        "Penalty includes expulsion from college plus legal action. "
        "Anti-Ragging Helpline: 1800-123-4567. "
        "National Anti-Ragging Helpline: 1800-180-5522. "
        "Complaint Email: antiragging@bvrithyderabad.edu.in, "
        "grievance@bvrithyderabad.edu.in, womenscell@bvrithyderabad.edu.in."
    )
    
    # Save
    doc.save("data/college_info.docx")
    print("✅ Created data/college_info.docx")


if __name__ == "__main__":
    generate_college_info()